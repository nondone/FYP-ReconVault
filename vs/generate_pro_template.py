# generate_pro_template.py
import os
from docx import Document

base = r"c:\Users\Vincent\Desktop\FYP Automated Recon\vs"
out = os.path.join(base, "templates", "pro_report_template.docx")

doc = Document()
doc.add_heading("ReconVault Professional Penetration Test Report", 0)

# ===== Report Metadata =====
doc.add_heading("1. Report Information", level=1)
meta = doc.add_table(rows=6, cols=2)
meta.style = "Table Grid"
meta.cell(0, 0).text = "Report ID"
meta.cell(0, 1).text = "{{ report_id }}"
meta.cell(1, 0).text = "Target"
meta.cell(1, 1).text = "{{ target }}"
meta.cell(2, 0).text = "Prepared For (User)"
meta.cell(2, 1).text = "{{ report_user }}"
meta.cell(3, 0).text = "Scanning Type"
meta.cell(3, 1).text = "{{ scan_type }}"
meta.cell(4, 0).text = "Scan Date"
meta.cell(4, 1).text = "{{ scan_date }}"
meta.cell(5, 0).text = "Generated At"
meta.cell(5, 1).text = "{{ generated_at }}"

# ===== Executive Summary =====
doc.add_heading("2. Executive Summary", level=1)
doc.add_paragraph("{{ exec_summary }}")

# ===== Scope =====
doc.add_heading("3. Scope", level=1)
doc.add_paragraph("In-scope assets for this assessment:")
scope_tbl = doc.add_table(rows=2, cols=3)
scope_tbl.style = "Table Grid"
scope_tbl.cell(0, 0).text = "Asset"
scope_tbl.cell(0, 1).text = "Type"
scope_tbl.cell(0, 2).text = "Notes"
scope_tbl.cell(1, 0).text = "{% for s in scope_assets %}{{ s.asset }}"
scope_tbl.cell(1, 1).text = "{{ s.asset_type }}"
scope_tbl.cell(1, 2).text = "{{ s.note }}{% endfor %}"

# ===== Methodology =====
doc.add_heading("4. Methodology", level=1)
doc.add_paragraph("{% for step in methodology_steps %}- {{ step }}\n{% endfor %}")

# ===== Modules Used =====
doc.add_heading("5. Modules Executed", level=1)
doc.add_paragraph("{% for m in modules_used %}- {{ m }}\n{% endfor %}")

# ===== Findings by Severity =====
doc.add_heading("6. Findings by Severity", level=1)
sev = doc.add_table(rows=5, cols=2)
sev.style = "Table Grid"
sev.cell(0, 0).text = "Critical"
sev.cell(0, 1).text = "{{ severity_counts.critical }}"
sev.cell(1, 0).text = "High"
sev.cell(1, 1).text = "{{ severity_counts.high }}"
sev.cell(2, 0).text = "Medium"
sev.cell(2, 1).text = "{{ severity_counts.medium }}"
sev.cell(3, 0).text = "Low"
sev.cell(3, 1).text = "{{ severity_counts.low }}"
sev.cell(4, 0).text = "Info"
sev.cell(4, 1).text = "{{ severity_counts.info }}"

# ===== OWASP Mapping =====
doc.add_heading("7. OWASP Top 10 Mapping", level=1)
doc.add_paragraph("{% for o in owasp_items %}- {{ o.code }}: {{ o.comment }}\n{% else %}No OWASP mapping identified.\n{% endfor %}")

# ===== Remediation =====
doc.add_heading("8. Remediation Plan", level=1)
doc.add_paragraph("{% for item in remediation_plan %}- {{ item }}\n{% endfor %}")

# ===== Subdomains =====
doc.add_heading("9. Discovered Subdomains", level=1)
sub_tbl = doc.add_table(rows=2, cols=2)
sub_tbl.style = "Table Grid"
sub_tbl.cell(0, 0).text = "No."
sub_tbl.cell(0, 1).text = "Subdomain"
sub_tbl.cell(1, 0).text = "{% for s in subdomain_rows %}{{ loop.index }}"
sub_tbl.cell(1, 1).text = "{{ s }}{% endfor %}"

# ===== Open Ports =====
doc.add_heading("10. Open Ports", level=1)
ports_tbl = doc.add_table(rows=2, cols=4)
ports_tbl.style = "Table Grid"
ports_tbl.cell(0, 0).text = "Host"
ports_tbl.cell(0, 1).text = "Port"
ports_tbl.cell(0, 2).text = "Protocol"
ports_tbl.cell(0, 3).text = "Service"
ports_tbl.cell(1, 0).text = "{% for p in open_ports %}{{ p.host }}"
ports_tbl.cell(1, 1).text = "{{ p.port }}"
ports_tbl.cell(1, 2).text = "{{ p.proto }}"
ports_tbl.cell(1, 3).text = "{{ p.service }}{% endfor %}"

# ===== Web Services =====
doc.add_heading("11. Web Services", level=1)
web_tbl = doc.add_table(rows=2, cols=3)
web_tbl.style = "Table Grid"
web_tbl.cell(0, 0).text = "URL"
web_tbl.cell(0, 1).text = "Status"
web_tbl.cell(0, 2).text = "Tech/Title"
web_tbl.cell(1, 0).text = "{% for w in web_services %}{{ w.url }}"
web_tbl.cell(1, 1).text = "{{ w.status }}"
web_tbl.cell(1, 2).text = "{{ w.detail }}{% endfor %}"

# ===== Appendix =====
doc.add_heading("12. Appendix", level=1)
doc.add_paragraph("Tools Used: {{ tools_used }}")
doc.add_paragraph("Raw vulnerability output is retained in system records.")

os.makedirs(os.path.dirname(out), exist_ok=True)
doc.save(out)
print("Template generated:", out)
